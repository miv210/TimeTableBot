import urllib.request
import win32com.client
import pythoncom
from bs4 import BeautifulSoup as bs
import requests
import docx
from library import bd
import os
import glob

class parser():
    def connect_site(self):
        headers = {"accept": "/",
                   "user-agent": 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/76.0.3809.132 Safari/537.36'}
        # Присоединение к главной странице сайта
        request = requests.get('http://portal.volgetc.ru/', headers=headers)
        soup = bs(request.content, 'html.parser')
        block_news = soup.findAll('ul', class_="latestnews")[1].find('a')['href']

        # Присоединение к странице с заменами
        request_timetable = requests.get('http://portal.volgetc.ru/' + str(block_news), headers=headers)
        soup = bs(request_timetable.content, 'html.parser')
        self.link = soup.find('iframe',class_="edocs_iframe")['src'][33:86]
        print(str(self.link))

    def doc_work(self,dir_file):
        file = urllib.request.urlopen(self.link).read()



        # Создание .doc фаила и запись в него
        doc_file = open('{0}/tabel.doc'.format(dir_file), 'wb')
        doc_file.write(file)
        doc_file.close()

        #Конвертация .doc в .docx
        word = win32com.client.Dispatch("Word.Application")


        for py in glob.iglob('{0}/tabel.doc'.format(dir_file)):
            in_file = os.path.abspath('{0}'.format(py))
            print(in_file)
            wb = word.Documents.Open(in_file)
            out_file = os.path.abspath("{0}x".format(py))
            print(out_file)
            wb.SaveAs2(out_file, FileFormat=16)  # file format for docx
            wb.Close()

        word.Quit()

    def pars(self, dir_file):
        doc = docx.Document('{0}/tabel.docx'.format(dir_file))

        # Подключени к бд
        conn = bd.connect_bd()
        query = ("""DELETE FROM timetable""")
        bd.update_bd(query, conn)

        #Парсим число и Заголовок
        paragraphs1 = ''
        for i in range(0, 3):
            paragraphs1 = paragraphs1 + '\n' + doc.paragraphs[i].text
        print(paragraphs1)

        #Парсим замены
        id = 0
        for table in doc.tables:
            amount_rows = len(table.rows)
            i = 0
            print(len(doc.tables))

            while i < amount_rows:
                try:
                    id += 1

                    section = table.cell(i, 0)
                    num_lesson = table.cell(i, 1)
                    subject = table.cell(i, 2)
                    teacher = table.cell(i, 3)
                    num_class = table.cell(i, 4)

                    #Обновляем данные в БД
                    query = ("""INSERT INTO timetable (id, name_group, pair_number, subject, teacher_fullname, room_number, date) 
        	                            VALUES 
        		                            ({0},'{1}','{2}', '{3}', '{4}', '{5}', '{6}')
                                     """).format(id, section.text, num_lesson.text, subject.text, teacher.text,
                                                 num_class.text, paragraphs1)
                    bd.update_bd(query, conn)

                except:
                    pass
                i += 1
        conn.close()

    def start(self, dir_file):
        self.connect_site()
        self.doc_work(dir_file)
        self.pars(dir_file)
